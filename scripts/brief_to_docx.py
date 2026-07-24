#!/usr/bin/env python3
"""Render a daily-brief markdown file (briefs/YYYY-MM-DD.md) into a matching
'[YYYY-MM-DD] Daily Brief.docx', styled in Inter Tight.

Usage: python3 scripts/brief_to_docx.py briefs/2026-07-23.md
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Inter Tight"

TIER_COLOR = {
    "MUST-KNOW": RGBColor(0xB0, 0x00, 0x20),
    "HIGH-IMPACT": RGBColor(0x1A, 0x5C, 0x2A),
    "SIGNAL": RGBColor(0x60, 0x60, 0x60),
}
MUTED = RGBColor(0x70, 0x70, 0x70)
LINK_BLUE = "1155CC"

TOKEN_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|_([^_]+)_")


def set_font(doc, style_name, size=None):
    st = doc.styles[style_name]
    st.font.name = FONT
    rpr = st.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if size:
        st.font.size = Pt(size)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich_line(paragraph, text, muted=False, italic=False):
    """Write text into paragraph, turning [label](url) into live hyperlinks,
    **bold** into bold runs, and _italic_ into italic runs (recursing so a
    link nested inside _..._ still becomes clickable)."""
    pos = 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            if muted or italic:
                run.italic = True
            if muted:
                run.font.color.rgb = MUTED
        if m.group(2):  # [label](url)
            add_hyperlink(paragraph, m.group(2), m.group(1))
        elif m.group(3) is not None:  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
            if italic:
                run.italic = True
            if muted:
                run.font.color.rgb = MUTED
        else:  # _italic_ (group 4) — recurse so nested links/bold still work
            add_rich_line(paragraph, m.group(4), muted=muted, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if muted or italic:
            run.italic = True
        if muted:
            run.font.color.rgb = MUTED


def parse_brief(md_text):
    lines = md_text.splitlines()
    title = lines[0].lstrip("# ").strip() if lines else "Daily Brief"

    sections = []  # list of (heading, body_lines)
    current_heading = None
    current_lines = []
    for line in lines[1:]:
        m = re.match(r"^##\s+(.*)", line)
        if m:
            if current_heading is not None:
                sections.append((current_heading, current_lines))
            current_heading = m.group(1).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections.append((current_heading, current_lines))
    return title, sections


def render_item_block(doc, block_lines):
    """block_lines: consecutive lines starting with '**Title**  ·  `TIER`' followed by
    '- **Source:** ...', '- **What:** ...', '- **Why us:** ...'"""
    text = "\n".join(block_lines).strip()
    if not text:
        return
    header_match = re.match(r"\*\*(.+?)\*\*\s*·\s*`([^`]+)`", text)
    if not header_match:
        # Not a structured item (e.g. free text like "Nothing cleared the bar") — just write as paragraph
        p = doc.add_paragraph()
        add_rich_line(p, text)
        return
    title, tier = header_match.group(1), header_match.group(2)
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    tag = p.add_run(f"   {tier}")
    tag.bold = True
    tag.font.size = Pt(9)
    tag.font.color.rgb = TIER_COLOR.get(tier, MUTED)

    for field_label, pattern in (
        ("Source", r"-\s*\*\*Source:\*\*\s*(.+)"),
        ("What", r"-\s*\*\*What:\*\*\s*(.+)"),
        ("Why us", r"-\s*\*\*Why us:\*\*\s*(.+)"),
    ):
        fm = re.search(pattern, text)
        if fm:
            fp = doc.add_paragraph()
            fp.add_run(f"{field_label}: ").bold = True
            add_rich_line(fp, fm.group(1).strip())
    doc.add_paragraph()


def build_docx(md_path: Path, out_path: Path):
    md_text = md_path.read_text(encoding="utf-8")
    title, sections = parse_brief(md_text)

    doc = Document()
    set_font(doc, "Normal", size=11)
    for style, size in [("Title", 26), ("Heading 1", 16), ("Heading 2", 13)]:
        set_font(doc, style, size=size)
    for style in ("List Number", "List Bullet"):
        set_font(doc, style, size=11)

    doc.add_heading(title, level=0)

    for heading, body_lines in sections:
        is_coverage_notes = heading.strip().lower() == "coverage notes"
        doc.add_heading(heading, level=2 if is_coverage_notes else 1)

        body_text = "\n".join(body_lines)

        if heading.strip().lower() == "must-knows":
            for m in re.finditer(r"^\d+\.\s+(.*)$", body_text, re.MULTILINE):
                p = doc.add_paragraph(style="List Number")
                add_rich_line(p, m.group(1).strip())
            continue

        if is_coverage_notes:
            for raw_line in body_lines:
                line = raw_line.strip()
                if not line or line == "---":
                    continue
                bullet = re.sub(r"^-\s*", "", line)
                p = doc.add_paragraph(style="List Bullet")
                add_rich_line(p, bullet, muted=True)
            continue

        # Regular section: split into item blocks on blank-line-separated chunks
        # that start with a **bold** header line.
        chunk = []
        for raw_line in body_lines + [""]:
            line = raw_line.rstrip()
            if line.strip() == "---":
                continue
            if re.match(r"^\*\*.+\*\*\s*·\s*`", line) and chunk:
                render_item_block(doc, chunk)
                chunk = [line]
            elif not line.strip() and chunk and not any(c.strip() for c in chunk[-1:]):
                continue
            else:
                chunk.append(line)
        if any(c.strip() for c in chunk):
            render_item_block(doc, chunk)

    # Footer (last line, e.g. "*Logged to ledger: ...*")
    footer_match = re.search(r"^\*Logged to ledger:.*\*$", md_text, re.MULTILINE)
    if footer_match:
        fp = doc.add_paragraph()
        r = fp.add_run(footer_match.group(0).strip("*"))
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 scripts/brief_to_docx.py briefs/YYYY-MM-DD.md", file=sys.stderr)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    date_match = re.search(r"(\d{4}-\d{2}-\d{2})", md_path.name)
    date_str = date_match.group(1) if date_match else md_path.stem
    out_path = md_path.parent / f"[{date_str}] Daily Brief.docx"
    build_docx(md_path, out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
